"""
Document parser for PDF, DOCX, and TXT files.
"""
import os
from pathlib import Path
from typing import Optional


class DocumentParser:
    """Parse documents and extract raw text."""

    @staticmethod
    def parse(file_path: str) -> str:
        """
        Parse a document and return raw text.
        Supports PDF, DOCX, and TXT formats.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()

        if file_ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif file_ext == ".docx":
            return DocumentParser._parse_docx(file_path)
        elif file_ext == ".txt":
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parse PDF using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed. Install with: pip install pdfplumber")

        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

        return "\n\n".join(text)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parse DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        doc = Document(file_path)
        text = []

        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text.append(row_text)

        return "\n".join(text)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Parse TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
